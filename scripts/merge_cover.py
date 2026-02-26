"""
merge_cover.py - 将封面+目录前缀模板与 Pandoc 生成的正文合并

用法:
    python merge_cover.py <模板docx> <正文docx> <输出docx> [--title "标题"]

依赖:
    pip install docxcompose
"""

import sys
import argparse
from docxcompose.composer import Composer
from docx import Document


def replace_title_placeholder(doc, title):
    """替换模板中的 {{TITLE}} 占位符为实际标题。
    
    WHY: Word 可能将 {{TITLE}} 拆分到多个 run 中，
    因此先尝试 run 级替换，若失败则回退到段落级重组。
    """
    PLACEHOLDER = '{{TITLE}}'
    
    for para in doc.paragraphs:
        if PLACEHOLDER not in para.text:
            continue
        
        # 尝试 run 级精确替换
        replaced = False
        for run in para.runs:
            if PLACEHOLDER in run.text:
                run.text = run.text.replace(PLACEHOLDER, title)
                replaced = True
        
        if replaced:
            continue
        
        # 回退：占位符被拆分到多个 run，用段落级重组
        # WHY: 保留第一个 run 的格式，清空其余 run
        new_text = para.text.replace(PLACEHOLDER, title)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''


def add_table_borders(doc):
    """给文档中所有表格添加全框线（单实线）。
    
    WHY: Pandoc 生成的 docx 表格默认无边框，且不使用模板的 Table Normal 样式，
    只能通过后处理在 OOXML 层面添加 tblBorders 元素。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    # 6 条边框：上、下、左、右、内部横线、内部竖线
    BORDER_NAMES = ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']
    # 边框属性：单实线，0.5pt（4 half-points），黑色
    BORDER_ATTRS = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}

    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, qn('w:tblPr'))

        # 移除已有的 tblBorders（避免重复）
        existing = tbl_pr.find(qn('w:tblBorders'))
        if existing is not None:
            tbl_pr.remove(existing)

        # 创建新的 tblBorders
        borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
        for name in BORDER_NAMES:
            border = etree.SubElement(borders, qn(f'w:{name}'))
            for attr, value in BORDER_ATTRS.items():
                border.set(qn(f'w:{attr}'), value)


def merge(prefix_path, body_path, output_path, title=None):
    """合并前缀模板与正文文档。
    
    Args:
        prefix_path: 包含封面+目录的模板文件路径
        body_path: Pandoc 生成的正文文件路径
        output_path: 合并后的输出文件路径
        title: 可选，替换 {{TITLE}} 占位符的标题文本
    """
    # 1. 打开前缀模板
    prefix = Document(prefix_path)
    
    # 2. 替换标题占位符
    if title:
        replace_title_placeholder(prefix, title)
    
    # 3. 给正文表格添加框线（在合并前处理，封面表格不受影响）
    body = Document(body_path)
    add_table_borders(body)
    
    # 4. 合并正文
    composer = Composer(prefix)
    composer.append(body)
    
    # 5. 保存
    composer.save(output_path)
    print(f"   合并完成: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='合并封面模板与正文文档')
    parser.add_argument('prefix', help='封面+目录模板文件路径')
    parser.add_argument('body', help='正文文档路径')
    parser.add_argument('output', help='输出文件路径')
    parser.add_argument('--title', help='替换 {{TITLE}} 占位符的标题', default=None)
    
    args = parser.parse_args()
    merge(args.prefix, args.body, args.output, args.title)


if __name__ == '__main__':
    main()
