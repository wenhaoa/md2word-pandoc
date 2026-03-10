"""
add_captions.py - Pandoc docx 后处理：题注域编号 + 表格/图片格式

用法:
    python add_captions.py <docx文件>

功能:
    1. 扫描 Pandoc 生成的 docx，识别 "图N-M" / "表N-M" 题注并替换为域
    2. 表格单元格垂直居中
    3. 图片段落居中

域代码结构:
    图{STYLEREF "Heading 1" \\s}-{SEQ 图 \\* ARABIC \\s 1} 标题

依赖:
    pip install python-docx lxml
"""

import sys
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# 匹配独立题注行
CAPTION_PATTERN = re.compile(r'^(图|表)\s*\d+[-\.]\d+\s*')
MAX_CAPTION_LENGTH = 60


def make_field_runs(instr_text):
    """创建域代码 run 序列（begin + instrText + separate + result + end）。

    Args:
        instr_text: 域指令文本，如 ' SEQ 图 \\* ARABIC \\s 1 '
    """
    runs = []

    # begin
    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)
    runs.append(r_begin)

    # instrText
    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    r_instr.append(instr)
    runs.append(r_instr)

    # separate
    r_sep = OxmlElement('w:r')
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fld_sep)
    runs.append(r_sep)

    # result placeholder
    r_result = OxmlElement('w:r')
    t_result = OxmlElement('w:t')
    t_result.text = '?'
    r_result.append(t_result)
    runs.append(r_result)

    # end
    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)
    runs.append(r_end)

    return runs


def apply_rpr_to_runs(run_elems, rpr_source_elem):
    """给 run 元素列表应用格式属性。"""
    if rpr_source_elem is None:
        return
    source_rpr = rpr_source_elem.find(qn('w:rPr'))
    if source_rpr is None:
        return
    for run_elem in run_elems:
        if run_elem.tag == qn('w:r'):
            new_rpr = copy.deepcopy(source_rpr)
            run_elem.insert(0, new_rpr)


def make_text_run(text):
    """创建文本 run。"""
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _ensure_bold(run_elems):
    """给 run 元素列表强制添加加粗格式。
    
    WHY: 题注（图N-M / 表N-M）需要加粗以区别于正文。
    """
    for r in run_elems:
        if r.tag != qn('w:r'):
            continue
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            r.insert(0, rpr)
        if rpr.find(qn('w:b')) is None:
            rpr.append(OxmlElement('w:b'))


def build_caption_runs(caption_type, rpr_source=None):
    """构建全域化题注 run 序列。

    生成: 图{STYLEREF "Heading 1" \\s}-{SEQ 图 \\* ARABIC \\s 1} 标题
    """
    all_runs = []

    # "图" / "表" 前缀
    all_runs.append(make_text_run(caption_type))

    # STYLEREF 域: 章节号
    # WHY: 数字 1 引用 Heading 1 样式级别, \\s 返回段落编号
    styleref_runs = make_field_runs(' STYLEREF 1 \\s ')
    all_runs.extend(styleref_runs)

    # 连字符
    all_runs.append(make_text_run('-'))

    # SEQ 域: 章内自动编号
    # WHY: \\s 1 在每个 Heading 1 处重新计数
    seq_runs = make_field_runs(f' SEQ {caption_type} \\* ARABIC \\s 1 ')
    all_runs.extend(seq_runs)

    # 空格
    all_runs.append(make_text_run(' '))

    if rpr_source is not None:
        apply_rpr_to_runs(all_runs, rpr_source)

    return all_runs


def set_caption_format(para):
    """设置题注段落格式：居中、无首行缩进。"""
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def has_drawing(para_elem):
    """检查段落是否包含图片。"""
    return bool(para_elem.findall('.//' + qn('w:drawing')))


def center_table_cells(doc):
    """设置所有表格单元格垂直居中。

    WHY: Pandoc 生成的表格默认靠上对齐，多行文字时视觉效果差。
    """
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # 移除已有的 vAlign（避免重复）
                for old in tcPr.findall(qn('w:vAlign')):
                    tcPr.remove(old)
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                count += 1
    if count > 0:
        print(f'  {count} table cell(s) vertically centered')


def center_images(doc):
    """将包含图片的段落设为居中对齐、无首行缩进。

    WHY: Pandoc 生成的图片段落继承 Normal 样式（左对齐+首行缩进），
    图片应居中显示。
    """
    count = 0
    for p in doc.paragraphs:
        if has_drawing(p._element):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            count += 1
    if count > 0:
        print(f'  {count} image paragraph(s) centered')


def process_captions(doc):
    """处理所有图表题注。"""
    count = 0
    paragraphs = doc.paragraphs

    for i in range(len(paragraphs)):
        p = paragraphs[i]
        text = p.text.strip()

        match = CAPTION_PATTERN.match(text)
        if not match:
            continue

        if len(text) > MAX_CAPTION_LENGTH:
            continue

        caption_type = match.group(1)

        # 图片题注：上一段落应含图片
        if caption_type == '图' and i > 0:
            prev_p = paragraphs[i - 1]
            if not has_drawing(prev_p._element) and not has_drawing(p._element):
                continue

        # 提取标题文本
        title_text = CAPTION_PATTERN.sub('', text)

        # 保存格式
        first_run_elem = p.runs[0]._element if p.runs else None

        # 清空 runs
        for r in p._element.findall(qn('w:r')):
            p._element.remove(r)

        # 构建域 runs（全域化：STYLEREF + SEQ）
        caption_runs = build_caption_runs(caption_type, first_run_elem)

        # 标题文本 run
        r_title = make_text_run(title_text)
        if first_run_elem is not None:
            source_rpr = first_run_elem.find(qn('w:rPr'))
            if source_rpr is not None:
                r_title.insert(0, copy.deepcopy(source_rpr))

        # 加粗全部题注 run（域编号 + 标题文本）
        all_runs = caption_runs + [r_title]
        _ensure_bold(all_runs)

        # 插入 runs
        ppr = p._element.find(qn('w:pPr'))
        insert_after = ppr if ppr is not None else None

        for run_elem in all_runs:
            if insert_after is not None:
                insert_after.addnext(run_elem)
                insert_after = run_elem
            else:
                p._element.append(run_elem)

        # 居中、无缩进
        set_caption_format(p)

        count += 1
        print(f'  ok {caption_type}: [{i}] "{caption_type}?-?  {title_text[:30]}"')

    return count


def main():
    if len(sys.argv) < 2:
        print('Usage: python add_captions.py <docx>')
        sys.exit(1)

    docx_path = sys.argv[1]
    print(f'  Processing: {docx_path}')

    doc = Document(docx_path)

    # 处理题注
    count = process_captions(doc)

    # 表格单元格垂直居中
    center_table_cells(doc)

    # 图片段落居中
    center_images(doc)

    doc.save(docx_path)
    print(f'  {count} caption(s) processed')
    print(f'  Tip: In Word press Ctrl+A then F9 to update fields')


if __name__ == '__main__':
    main()
